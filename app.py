# -*- coding: utf-8 -*-
"""
DSE/高考数学提分潜力诊断工具 v2.1
陈老师专属 - AI驱动的数学诊断系统
参考 Math AI Insight Pro 标准
"""

import streamlit as st
import google.generativeai as genai
from zhipuai import ZhipuAI
from PIL import Image, ImageEnhance
import fitz  # PyMuPDF
import io
import re
import json
import time
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING

# ==================== 核心配置 ====================
GEMINI_API_KEY = "AIzaSyBcvLsNA4ZeLbxHjcWmx_Fy1OcXYS5z9J0"
GLM_API_KEY = "445b29b7119946d49c65361161dae089.tdSIhpAFssxWAoEO"
WECHAT_ID = "xiaobo20230512"

PRIMARY_MODEL = "gemini-2.5-pro"
FALLBACK_MODEL = "glm-4-plus"

# ==================== UI 配置 ====================
st.set_page_config(
    page_title="陈老师数学诊断",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
<style>
    /* 全局深色动态背景 */
    @keyframes gradient-bg {
        0% {background-position: 0% 50%;}
        50% {background-position: 100% 50%;}
        100% {background-position: 0% 50%;}
    }
    .stApp {
        background: linear-gradient(-45deg, #0b0f19, #1b2735, #243b55, #141e30);
        background-size: 400% 400%;
        animation: gradient-bg 15s ease infinite;
        color: #e0e0e0;
    }
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}

    /* 动画效果 */
    @keyframes gradient-text {
        0% { background-position: 0% 50%; }
        50% { background-position: 100% 50%; }
        100% { background-position: 0% 50%; }
    }
    @keyframes buttonBreath {
        0% { box-shadow: 0 0 5px rgba(0, 201, 255, 0.3); transform: scale(1); }
        50% { box-shadow: 0 0 20px rgba(0, 201, 255, 0.6); transform: scale(1.02); }
        100% { box-shadow: 0 0 5px rgba(0, 201, 255, 0.3); transform: scale(1); }
    }

    /* Tab 按钮样式 */
    .stTabs [data-baseweb="tab-list"] { gap: 8px; }
    .stTabs [data-baseweb="tab"] {
        height: 45px;
        background-color: rgba(255, 255, 255, 0.05);
        border-radius: 6px;
        border: 1px solid rgba(255,255,255,0.1);
        color: #aaa;
        transition: all 0.3s;
        font-size: 14px;
    }
    .stTabs [aria-selected="true"] {
        background: linear-gradient(90deg, #00C9FF, #92FE9D);
        color: #000 !important;
        font-weight: 700;
        border: none;
        box-shadow: 0 0 10px rgba(0, 201, 255, 0.3);
    }

    /* 侧边栏样式 - 全面白色高亮 */
    [data-testid="stSidebar"] {
        background-color: #050505;
        border-right: 1px solid rgba(255, 255, 255, 0.1);
    }
    [data-testid="stSidebar"] * {
        color: #FFFFFF !important;
        -webkit-text-fill-color: #FFFFFF !important;
    }
    [data-testid="stSidebar"] h1,
    [data-testid="stSidebar"] h2,
    [data-testid="stSidebar"] h3,
    [data-testid="stSidebar"] h4,
    [data-testid="stSidebar"] h5,
    [data-testid="stSidebar"] h6 {
        color: #FFFFFF !important;
        -webkit-text-fill-color: #FFFFFF !important;
    }
    [data-testid="stSidebar"] label {
        color: #FFFFFF !important;
        font-size: 0.95rem !important;
        font-weight: 600 !important;
    }
    [data-testid="stSidebar"] p,
    [data-testid="stSidebar"] span,
    [data-testid="stSidebar"] div {
        color: #E0E0E0 !important;
    }
    [data-testid="stSidebar"] [data-baseweb="input"] > div,
    [data-testid="stSidebar"] [data-baseweb="select"] > div {
        background-color: #1a1a1a !important;
        border: 1px solid #4a9eff !important;
        color: #FFFFFF !important;
        border-radius: 8px !important;
    }
    [data-testid="stSidebar"] input,
    [data-testid="stSidebar"] textarea {
        background-color: #1a1a1a !important;
        color: #FFFFFF !important;
        -webkit-text-fill-color: #FFFFFF !important;
    }
    [data-testid="stSidebar"] [role="option"] {
        background-color: #1a1a1a !important;
        color: #FFFFFF !important;
    }
    [data-testid="stSidebar"] [data-baseweb="select"] span {
        color: #FFFFFF !important;
    }
    [data-testid="stSidebar"] .stSelectbox label {
        color: #FFFFFF !important;
    }

    /* 上传框样式 */
    [data-testid='stFileUploader'] * {
        color: #FFFFFF !important;
        -webkit-text-fill-color: #FFFFFF !important;
        opacity: 1 !important;
    }
    [data-testid='stUploadedFileItem'] {
        background-color: rgba(255, 255, 255, 0.1) !important;
        border: 1px solid #00C9FF !important;
        border-radius: 8px !important;
        padding: 10px !important;
    }
    [data-testid='stFileUploader'] button,
    [data-testid='stFileUploader'] button * {
        color: #000000 !important;
        -webkit-text-fill-color: #000000 !important;
        font-weight: 800 !important;
    }
    [data-testid='stFileUploader'] [data-testid='stUploadedFileItem'] button,
    [data-testid='stFileUploader'] [data-testid='stUploadedFileItem'] button * {
        color: #FF4B4B !important;
        -webkit-text-fill-color: #FF4B4B !important;
        fill: #FF4B4B !important;
    }
    [data-testid='stFileUploader'] label {
        color: #FFD700 !important;
        -webkit-text-fill-color: #FFD700 !important;
        font-size: 1.1rem !important;
        font-weight: 800 !important;
        text-shadow: 0 0 8px rgba(255, 215, 0, 0.4) !important;
    }
    [data-testid='stFileUploader'] button {
        background: linear-gradient(90deg, #00C9FF, #5EE7DF, #92FE9D, #00C9FF) !important;
        background-size: 300% 100% !important;
        border: none !important;
        border-radius: 20px !important;
        padding: 8px 20px !important;
        animation: gradient-text 4s linear infinite, buttonBreath 3s ease-in-out infinite !important;
        box-shadow: 0 5px 15px rgba(0, 201, 255, 0.4);
    }
    [data-testid='stFileUploader'] button:hover {
        transform: scale(1.03) translateY(-2px) !important;
        box-shadow: 0 10px 25px rgba(0, 201, 255, 0.6) !important;
    }
    [data-testid='stFileUploader'] section {
        background-color: rgba(30, 34, 45, 0.6);
        border: 1px dashed rgba(0, 201, 255, 0.5) !important;
        border-radius: 10px;
        padding: 25px 20px !important;
        min-height: 125px;
    }

    /* 全局按钮样式 */
    .stButton > button, .stDownloadButton > button {
        background: linear-gradient(90deg, #00C9FF, #5EE7DF, #92FE9D, #00C9FF) !important;
        background-size: 300% 100% !important;
        color: #000 !important;
        border: none !important;
        border-radius: 50px !important;
        font-weight: 800 !important;
        animation: gradient-text 4s linear infinite, buttonBreath 3s ease-in-out infinite !important;
        box-shadow: 0 5px 15px rgba(0, 201, 255, 0.4);
        transition: all 0.3s;
    }
    .stButton > button:hover, .stDownloadButton > button:hover {
        transform: scale(1.03) translateY(-2px) !important;
        box-shadow: 0 10px 25px rgba(0, 201, 255, 0.6) !important;
    }

    /* 标题样式 */
    h1 {
        font-family: 'Segoe UI', sans-serif;
        font-weight: 800;
        background: linear-gradient(90deg, #00C9FF 0%, #92FE9D 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        text-align: center;
        font-size: 2.2rem;
    }
    .glass-card {
        background: rgba(255, 255, 255, 0.03);
        backdrop-filter: blur(16px);
        border: 1px solid rgba(255, 255, 255, 0.08);
        border-radius: 16px;
        padding: 20px;
        margin-bottom: 20px;
    }

    /* 提示框样式 */
    [data-testid="stAlert"] > div, [data-testid="stAlert"] p {
        color: #FFFFFF !important;
        -webkit-text-fill-color: #FFFFFF !important;
    }
    [data-testid="stAlert"] svg {
        fill: #FFFFFF !important;
    }

    /* ========== 主内容区标题高亮（参考 Math AI Pro 标准）========== */
    /* 全局标签样式修复 - 确保在深色背景下可见 */
    label, .label, [role="label"], span[kind="label"] {
        color: #FFFFFF !important;
        -webkit-text-fill-color: #FFFFFF !important;
    }

    /* 所有 Streamlit 输入组件的标签 */
    [data-testid*="stNumberInput"] label,
    [data-testid*="stSelect"] label,
    [data-testid*="stMulti"] label,
    [data-testid*="stRadio"] label,
    [data-testid*="stSlider"] label,
    [data-testid*="stText"] label,
    [data-testid*="stFileUpload"] label {
        color: #00C9FF !important;
        -webkit-text-fill-color: #00C9FF !important;
        font-weight: 700 !important;
    }

    /* 多选框下拉框内文本 */
    [data-testid="stMultiselect"] span {
        color: #FFFFFF !important;
        -webkit-text-fill-color: #FFFFFF !important;
    }

    /* 下拉框选项 */
    [data-testid="stSelectbox"] option {
        background-color: #1a1a1a !important;
        color: #FFFFFF !important;
    }

    /* 单选框选项 */
    [data-testid="stRadio"] [role="radio"] {
        background-color: #1a1a1a !important;
        border-color: #00C9FF !important;
    }
    [data-testid="stRadio"] [role="radio"] + div {
        color: #00C9FF !important;
    }

    /* 所有文本输入确保可见 */
    input[type="text"], input[type="number"], textarea {
        background-color: #1a1a1a !important;
        color: #FFFFFF !important;
        -webkit-text-fill-color: #FFFFFF !important;
    }
</style>
""", unsafe_allow_html=True)

# ==================== AI 调用函数 ====================
def call_ai_gemini(prompt, images=None):
    """使用Gemini API"""
    try:
        genai.configure(api_key=GEMINI_API_KEY)
        model = genai.GenerativeModel(PRIMARY_MODEL)
        if images:
            response = model.generate_content([prompt] + images)
        else:
            response = model.generate_content(prompt)
        return response.text, "gemini"
    except Exception as e:
        error_msg = str(e)
        # 详细错误信息
        if "401" in error_msg or "UNAUTHENTICATED" in error_msg:
            return None, "Gemini API密钥无效，请检查配置"
        elif "429" in error_msg:
            return None, "Gemini API请求过于频繁，请稍后重试"
        elif "quota" in error_msg.lower():
            return None, "Gemini API配额已用完"
        else:
            return None, f"Gemini错误: {error_msg[:100]}"

def call_ai_glm(prompt):
    """使用GLM API作为备用"""
    try:
        client = ZhipuAI(api_key=GLM_API_KEY)
        response = client.chat.completions.create(
            model=FALLBACK_MODEL,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=4000,
            temperature=0.7
        )
        return response.choices[0].message.content, "glm"
    except Exception as e:
        error_msg = str(e)
        if "401" in error_msg or "UNAUTHORIZED" in error_msg:
            return None, "GLM API密钥无效，请检查配置"
        elif "429" in error_msg:
            return None, "GLM API请求过于频繁，请稍后重试"
        else:
            return None, f"GLM错误: {error_msg[:100]}"

def call_ai_with_retry(model, prompt, content_list=None, max_retries=3, retry_delay=30):
    """带重试的AI调用"""
    for attempt in range(max_retries):
        try:
            if content_list:
                return model.generate_content([prompt] + content_list)
            else:
                return model.generate_content(prompt)
        except Exception as e:
            if "429" in str(e) and attempt < max_retries - 1:
                placeholder = st.empty()
                progress_text = f"⚠️ API正忙，自动排队中... (尝试 {attempt+1}/{max_retries})"
                my_bar = placeholder.progress(0, text=progress_text)
                for i in range(retry_delay):
                    time.sleep(1)
                    my_bar.progress((i+1)/retry_delay, text=f"⏳ 剩余 {retry_delay-i}s")
                placeholder.empty()
                continue
            else:
                raise e

def call_ai_with_fallback(prompt, images=None):
    """智能调用AI，自动切换备用模型"""
    result, source = call_ai_gemini(prompt, images)
    if result:
        return result, source
    if images:
        return None, "图像输入需要Gemini，当前服务繁忙"
    result, source = call_ai_glm(prompt)
    if result:
        return result, source
    return None, "所有AI服务暂时不可用，请稍后重试"

# ==================== 辅助函数 ====================
def enhance_image_for_ocr(pil_image):
    """增强图像用于OCR识别"""
    enhancer = ImageEnhance.Contrast(pil_image)
    img = enhancer.enhance(1.5)
    enhancer = ImageEnhance.Sharpness(img)
    img = enhancer.enhance(1.5)
    return img

def process_pdf_bytes(file_bytes, start_page, end_page):
    """处理PDF文件，提取图像"""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    total_pages = len(doc)
    start = max(0, start_page - 1)
    end = min(total_pages, end_page)

    images = []
    enhanced_images = []

    for i in range(start, end):
        try:
            page = doc.load_page(i)
            pix = page.get_pixmap(dpi=200)
            img_data = pix.tobytes("png")
            original = Image.open(io.BytesIO(img_data))

            if original.width > 2000:
                ratio = 2000 / original.width
                new_size = (2000, int(original.height * ratio))
                original = original.resize(new_size, Image.Resampling.LANCZOS)

            images.append(original)
            enhanced_images.append(enhance_image_for_ocr(original))
            del pix, img_data
        except Exception as e:
            st.error(f"页码 {i+1} 处理错误: {e}")

    doc.close()
    return images, enhanced_images

def clean_markdown_text(text):
    """清理Markdown文本"""
    text = text.replace("---", "")
    return text

# ==================== 雷达图（参考 Math AI Insight Pro 标准 - Linux兼容）====================
def create_radar_chart_image(scores):
    """创建雷达图 - 白底专业版（跨平台兼容）"""
    labels = list(scores.keys())
    values = list(scores.values())
    values += values[:1]
    angles = np.linspace(0, 2 * np.pi, len(labels), endpoint=False).tolist()
    angles += angles[:1]

    fig, ax = plt.subplots(figsize=(7, 7), subplot_kw=dict(polar=True))
    fig.patch.set_facecolor('white')
    ax.set_facecolor('white')
    ax.grid(color='#E9E9E9', linestyle='-', linewidth=1.0)
    ax.spines['polar'].set_visible(False)

    ax.plot(angles, values, color='#0066CC', linewidth=2.5, linestyle='-', zorder=10)
    ax.fill(angles, values, color='#0066CC', alpha=0.15)
    ax.scatter(angles, values, color='#0066CC', s=80, edgecolors='white', linewidth=2, zorder=11)
    ax.set_ylim(0, 100)

    # 设置标签 - 不使用字体文件，直接设置
    ax.set_yticklabels([])
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(labels, color='black', weight='bold', fontsize=13)

    ax.tick_params(pad=35)

    # 保存图片
    img_buf = io.BytesIO()
    plt.savefig(img_buf, format='png', bbox_inches='tight', dpi=300, facecolor='white', transparent=False)
    img_buf.seek(0)
    plt.close(fig)
    return img_buf

# ==================== Word 排版（参考 Math AI Insight Pro 标准）====================
def set_font(run, font_name_cn, font_name_en='Times New Roman', size_pt=10.5, bold=False, italic=False, color=None):
    """设置Word字体 - 中英文分别设置"""
    run.font.name = font_name_en
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name_cn)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def add_page_number(run):
    """添加页码"""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def add_num_pages(run):
    """添加总页数"""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "NUMPAGES"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def clean_latex_for_word(text):
    """将LaTeX公式转换为Word可读的文本格式"""
    text = text.replace(r'\times', '×').replace(r'\cdot', '·')
    text = text.replace(r'\div', '÷')
    text = text.replace(r'\le', '≤').replace(r'\ge', '≥').replace(r'\neq', '≠')
    text = text.replace(r'\approx', '≈').replace(r'\%', '%')
    text = re.sub(r'\\frac\{(.*?)\}\{(.*?)\}', r'(\1)/(\2)', text)
    text = text.replace(r'\_', '_')
    text = re.sub(r'\\text\{(.*?)\}', r'\1', text)
    text = re.sub(r'\\mathbf\{(.*?)\}', r'\1', text)
    text = text.replace('$', '').replace('\\', '')
    return text

def create_word_docx_simple(report_text, student_name, radar_img_stream=None):
    """创建Word文档 - 麦肯锡报告风格"""
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # 设置默认样式
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(10.5)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # 颜色定义
    COLOR_CORRECT = RGBColor(0, 176, 80)
    COLOR_PARTIAL = RGBColor(237, 125, 49)
    COLOR_WRONG = RGBColor(255, 0, 0)
    COLOR_BLANK = RGBColor(128, 128, 128)
    COLOR_BLACK = RGBColor(0, 0, 0)
    COLOR_RED_HIGHLIGHT = RGBColor(255, 0, 0)

    # 页脚
    footer = section.footer
    p_footer = footer.paragraphs[0]
    p_footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_footer = p_footer.add_run("第 ")
    set_font(run_footer, '宋体', size_pt=9)
    add_page_number(p_footer.add_run())
    run_footer = p_footer.add_run(" 页 共 ")
    set_font(run_footer, '宋体', size_pt=9)
    add_num_pages(p_footer.add_run())
    run_footer = p_footer.add_run(" 页")
    set_font(run_footer, '宋体', size_pt=9)

    # 处理报告内容
    if "---JSON_START---" in report_text:
        main_content = report_text.split("---JSON_START---")[0]
    else:
        main_content = report_text

    main_content = clean_markdown_text(main_content)
    lines = main_content.split('\n')
    radar_inserted = False
    in_summary_section = False

    for line in lines:
        line = line.strip()
        if not line or line.startswith("```"):
            continue

        # 一级标题
        if line.startswith('# '):
            p = doc.add_heading(level=1)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run(line.replace('# ', ''))
            set_font(run, '黑体', 'Times New Roman', 18, bold=True, color=COLOR_BLACK)
            p.paragraph_format.space_after = Pt(20)

        # 二级标题
        elif line.startswith('## '):
            p = doc.add_heading(level=2)
            clean_text = line.replace('## ', '')
            run = p.add_run(clean_text)
            set_font(run, '微软雅黑', 'Microsoft YaHei', 15, bold=True, color=COLOR_BLACK)
            p.paragraph_format.space_before = Pt(12)

            if "总结" in clean_text or "展望" in clean_text:
                in_summary_section = True
            else:
                in_summary_section = False

            # 在"总体表现"后插入雷达图
            if ("总体" in clean_text or "概览" in clean_text) and not radar_inserted and radar_img_stream:
                radar_img_stream.seek(0)
                p_img = doc.add_paragraph()
                p_img.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                p_img.add_run().add_picture(radar_img_stream, width=Inches(4.5))
                radar_inserted = True

        # 三级标题
        elif line.startswith('### '):
            p = doc.add_heading(level=3)
            clean_text = line.replace('### ', '')
            title_color = COLOR_BLACK
            if "[完全正确]" in clean_text:
                title_color = COLOR_CORRECT
            elif "[部分正确]" in clean_text:
                title_color = COLOR_PARTIAL
            elif "[错误]" in clean_text:
                title_color = COLOR_WRONG
            elif "[未作答]" in clean_text:
                title_color = COLOR_BLANK
            run = p.add_run(clean_text)
            set_font(run, '黑体', 'Times New Roman', 12, bold=True, color=title_color)

        # 普通段落和列表
        else:
            if line.startswith('* ') or line.startswith('- '):
                p = doc.add_paragraph(style='List Bullet')
                clean_line = line[2:]
            else:
                p = doc.add_paragraph()
                clean_line = line

            is_highlight_line = "试卷得分" in clean_line or "正确率" in clean_line
            if in_summary_section:
                is_highlight_line = False

            # 处理格式
            parts = re.split(r'(\*\*.*?\*\*|`.*?`)', clean_line)
            for part in parts:
                if not part:
                    continue
                if part.startswith('**') and part.endswith('**'):
                    text = part[2:-2]
                    run = p.add_run(text)
                    if is_highlight_line:
                        color = COLOR_RED_HIGHLIGHT
                    elif in_summary_section:
                        color = COLOR_BLACK
                    else:
                        is_score = re.search(r'\d+分|\d+%', text)
                        color = RGBColor(255, 0, 0) if is_score else COLOR_BLACK
                    set_font(run, '宋体', 'Times New Roman', 10.5, bold=True, color=color)
                elif part.startswith('`') and part.endswith('`'):
                    clean_math_text = clean_latex_for_word(part[1:-1])
                    run = p.add_run(clean_math_text)
                    color = COLOR_RED_HIGHLIGHT if is_highlight_line else None
                    if in_summary_section:
                        color = COLOR_BLACK
                    set_font(run, 'Times New Roman', 'Times New Roman', 10.5, bold=False, italic=True, color=color)
                else:
                    run = p.add_run(part)
                    color = COLOR_RED_HIGHLIGHT if is_highlight_line else None
                    if in_summary_section:
                        color = COLOR_BLACK
                    set_font(run, '宋体', 'Times New Roman', 10.5, color=color)

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

# ==================== Prompt 模板 ====================
def get_quick_diagnosis_prompt(student_info):
    """快速诊断Prompt"""
    return f"""你是陈老师，一位有11年经验的DSE/高考数学专家。

【学生信息】
{json.dumps(student_info, ensure_ascii=False, indent=2)}

【任务】
请生成一份简明的数学诊断报告（预览版），包含：

1. **当前水平评估**（1-2句话，客观具体）
2. **主要问题识别**（3个要点，针对错题类型）
3. **提分建议**（3条具体可执行的建议）
4. **能力雷达图评分**（JSON格式，6个维度各0-100分）

【输出格式】
# {student_info.get('name', '同学')} 数学诊断报告（预览版）

## 一、当前水平评估
[具体评估内容，结合成绩和错题分析]

## 二、主要问题识别
1. [针对第一个错题类型的问题分析]
2. [针对第二个错题类型的问题分析]
3. [针对第三个错题类型的问题分析]

## 三、提分建议
1. [第一条具体建议，包含方法和时间]
2. [第二条具体建议]
3. [第三条具体建议]

## 四、获取完整报告
这是预览版（30%内容）。完整版包含：
- 详细知识漏洞分析
- 个性化学习计划（分阶段）
- 专属练习题库
- 提分时间预测

添加陈老师微信免费领取完整报告：{WECHAT_ID}

---JSON_START---
{{"代数运算": 65, "几何直观": 60, "逻辑推理": 70, "数据分析": 55, "数学建模": 50, "创新意识": 60}}
"""

def get_deep_diagnosis_prompt(student_name, exam_type):
    """深度诊断Prompt"""
    import datetime
    today_str = datetime.date.today().strftime("%Y年%m月%d日")

    return f"""
你是一位严谨的数学名师。请基于试卷图片撰写深度分析报告。

【学生信息】
- 姓名：{student_name}
- 考试类型：{exam_type}
- 报告日期：{today_str}

【指令】：
1. 仔细分析试卷中的每一道题
2. 识别学生的核心考点掌握情况
3. 给出针对性的学习建议

【输出格式】
# 《{student_name}同学试卷分析报告》
学生姓名：{student_name}
考试科目：{exam_type}
报告撰写：{today_str}
撰写人：陈老师

## 一、总体表现概览
* **试卷得分**：[推断分数]
* **正确率**：[计算百分比]%
* **总体评价**：[简练客观]

## 二、逐题深度分析
（请遍历每一道题，给出分析）

### 第X题 [状态]
* **核心考点**：**[考点]**
* **诊断分析**：[2-3行分析]
* **易错点**：[内容]
* **复习建议**：[具体建议]

## 三、能力薄弱点诊断
...

## 四、巩固知识与优势识别
...

## 五、阶段性复习建议与行动方案
1. 基础夯实阶段（建议：2-3周）
   * ...
2. 能力提升阶段（建议：基础阶段后3-4周）
   * ...
3. 应试与策略优化
   * ...

## 六、总结与展望
[内容，包含争取正确率突破XX%的期望]

---JSON_START---
{{
    "代数运算": [基于表现打分],
    "几何直观": [打分],
    "逻辑推理": [打分],
    "数据分析": [打分],
    "数学建模": [打分],
    "创新意识": [打分]
}}
"""

# ==================== 侧边栏 ====================
with st.sidebar:
    st.markdown("### ⚙️ 诊断设置")

    exam_type = st.selectbox(
        "考试类型",
        ("DSE - 必修数学", "DSE - 延伸M1", "DSE - 延伸M2", "高考 - 数学"),
        label_visibility="visible"
    )

    student_name = st.text_input("学生姓名", value="同学", placeholder="请输入姓名")

    st.markdown("---")
    st.markdown(f"""
    <div style='background: rgba(74, 158, 255, 0.1); border: 1px solid #4a9eff; border-radius: 10px; padding: 15px;'>
        <h4 style='color: #ffffff; margin: 0 0 10px 0;'>关于陈老师</h4>
        <p style='color: #8892b0; margin: 5px 0;'>11年数学教学经验</p>
        <p style='color: #8892b0; margin: 5px 0;'>3年DSE国际教育经验</p>
        <p style='color: #8892b0; margin: 5px 0 15px 0;'>专业：DSE延伸数学</p>
        <p style='color: #ffffff; margin: 0;'>微信：<strong>{WECHAT_ID}</strong></p>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
    <div style='text-align: center; color: #E0E0E0; font-size: 0.8em; margin-top: 20px;'>
        v2.1 Pro Standard<br>
        Designed & Developed by Porter
    </div>
    """, unsafe_allow_html=True)

# ==================== 主内容区 ====================
st.title("🧬 DSE/高考数学提分潜力诊断")

# 模式选择
mode = st.radio(
    "选择诊断模式",
    ["快速诊断", "深度诊断"],
    horizontal=True,
    label_visibility="collapsed"
)

if 'mode' not in st.session_state:
    st.session_state['mode'] = 'quick'

current_mode = 'quick' if mode == "快速诊断" else 'deep'
if st.session_state['mode'] != current_mode:
    st.session_state['mode'] = current_mode

# ==================== 快速诊断模式 ====================
if mode == "快速诊断":
    st.markdown("""
    <div class="glass-card">
        <h3>快速诊断</h3>
        <p>填写基本信息，AI系统将快速分析学生的数学学习状况，识别薄弱环节，给出针对性建议</p>
    </div>
    """, unsafe_allow_html=True)

    col1, col2 = st.columns(2)
    with col1:
        recent_score = st.number_input("最近一次数学成绩", min_value=0, max_value=160, value=80, step=1)
    with col2:
        total_score = st.number_input("试卷满分", min_value=60, max_value=160, value=150, step=1)

    col1, col2 = st.columns(2)
    with col1:
        wrong_topics = st.multiselect(
            "常错题型（可多选）",
            ["函数与导数", "三角函数", "数列", "解析几何", "概率统计", "立体几何", "延伸数学-微积分", "延伸数学-代数"],
            default=[]
        )
    with col2:
        learning_goal = st.selectbox(
            "学习目标",
            ["夯实基础", "提升成绩", "冲刺高分", "DSE延伸数学入门"]
        )

    if st.button("开始AI诊断", type="primary", use_container_width=True):
        if not wrong_topics:
            st.error("请至少选择一个错题类型，以便AI进行精准分析")
        else:
            student_info = {
                "name": student_name,
                "score": recent_score,
                "total": total_score,
                "wrong_topics": wrong_topics,
                "goal": learning_goal,
                "exam_type": exam_type
            }

            with st.status("AI正在分析中...", expanded=True) as status:
                st.write("1. 分析成绩数据...")
                time.sleep(0.3)
                st.write("2. 识别薄弱环节...")
                time.sleep(0.3)
                st.write("3. 生成诊断报告...")

                prompt = get_quick_diagnosis_prompt(student_info)
                result, source = call_ai_with_fallback(prompt)

                if result:
                    if "---JSON_START---" in result:
                        parts = result.split("---JSON_START---")
                        body = parts[0].strip()
                        json_str = parts[1].strip().replace("```json", "").replace("```", "").strip()
                    else:
                        body = result
                        json_str = '{"代数运算": 60, "几何直观": 60, "逻辑推理": 60, "数据分析": 60, "数学建模": 60, "创新意识": 60}'

                    try:
                        radar_data = json.loads(json_str)
                    except:
                        radar_data = {"代数运算": 60, "几何直观": 60, "逻辑推理": 60, "数据分析": 60, "数学建模": 60, "创新意识": 60}

                    st.session_state['report_text'] = body
                    st.session_state['radar_img'] = create_radar_chart_image(radar_data)
                    st.session_state['student_name'] = student_name

                    st.toast(f"诊断完成！使用模型：{source}", icon="✅")
                    status.update(label="✅ 诊断完成！", state="complete")
                    st.rerun()
                else:
                    st.error(f"诊断失败：{source}")

    # 显示报告
    if 'report_text' in st.session_state:
        col1, col2 = st.columns([3, 2])

        with col1:
            st.markdown("<h3 style='color: #ffffff; margin-bottom: 15px;'>诊断报告</h3>", unsafe_allow_html=True)
            st.markdown(f"<div class='glass-card'>{st.session_state['report_text']}</div>", unsafe_allow_html=True)

        with col2:
            if 'radar_img' in st.session_state:
                st.image(st.session_state['radar_img'], use_container_width=True, caption="能力维度分析")

            docx_file = create_word_docx_simple(
                st.session_state['report_text'],
                st.session_state.get('student_name', '同学'),
                st.session_state.get('radar_img')
            )
            st.download_button(
                "📥 下载报告",
                data=docx_file,
                file_name=f"{st.session_state.get('student_name', '同学')}_诊断报告.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )

        # 微信引流
        st.markdown(f"""
        <div style='background: linear-gradient(90deg, rgba(0,201,255,0.2), rgba(146,254,157,0.2)); padding: 20px; border-radius: 10px; text-align: center; margin-top: 20px; border: 1px solid #00C9FF;'>
            <h3 style='color: #ffffff;'>🔥 获取完整深度报告</h3>
            <p style='color: #8892b0;'>完整版包含详细知识漏洞分析、个性化学习计划、专属练习题库</p>
            <p style='font-size: 1.2rem; font-weight: bold; margin: 15px 0; color: #00C9FF;'>微信：{WECHAT_ID}</p>
            <p style='color: #8892b0;'>备注【提分】免费领取完整报告</p>
        </div>
        """, unsafe_allow_html=True)

# ==================== 深度诊断模式 ====================
elif mode == "深度诊断":
    st.markdown("""
    <div class="glass-card">
        <h3>深度诊断</h3>
        <p>上传试卷图片或PDF，AI将逐题分析，生成详细的学习诊断报告</p>
    </div>
    """, unsafe_allow_html=True)

    uploaded_file = st.file_uploader(
        "📂 上传试卷图片或PDF",
        type=['pdf', 'png', 'jpg', 'jpeg'],
        help="支持PDF、PNG、JPG格式，建议上传清晰图片"
    )

    if uploaded_file:
        file_bytes = uploaded_file.getvalue()

        if uploaded_file.type == "application/pdf":
            doc_temp = fitz.open(stream=file_bytes, filetype="pdf")
            total_pages = len(doc_temp)
            doc_temp.close()

            st.info(f"📄 检测到PDF文件，共 {total_pages} 页")

            page_range = st.slider("选择要分析的页面", 1, total_pages, (1, min(3, total_pages)))

            if st.button("🚀 开始深度分析", type="primary", use_container_width=True):
                with st.status("🔍 AI分析中...", expanded=True) as status:
                    st.write("1. 处理图像...")
                    images, enhanced = process_pdf_bytes(file_bytes, page_range[0], page_range[1])

                    st.write("2. AI逐题扫描...")
                    try:
                        genai.configure(api_key=GEMINI_API_KEY)
                        model = genai.GenerativeModel(PRIMARY_MODEL)

                        prompt = get_deep_diagnosis_prompt(student_name, exam_type)
                        response = call_ai_with_retry(model, prompt, enhanced)
                        full_text = response.text

                        st.write("3. 生成雷达图...")

                        if "---JSON_START---" in full_text:
                            parts = full_text.split("---JSON_START---")
                            body = parts[0].strip()
                            json_str = parts[1].strip().replace("```json", "").replace("```", "").strip()
                        else:
                            body = full_text
                            json_str = '{"代数运算": 60, "几何直观": 60, "逻辑推理": 60, "数据分析": 60, "数学建模": 60, "创新意识": 60}'

                        try:
                            radar_data = json.loads(json_str)
                        except:
                            radar_data = {"代数运算": 60, "几何直观": 60, "逻辑推理": 60, "数据分析": 60, "数学建模": 60, "创新意识": 60}

                        st.session_state['report_text'] = body
                        st.session_state['radar_img'] = create_radar_chart_image(radar_data)
                        st.session_state['student_name'] = student_name

                        st.toast("✅ 分析完成！", icon="🎉")
                        status.update(label="✅ 分析完成！", state="complete")
                        st.rerun()

                    except Exception as e:
                        st.error(f"分析失败：{e}")

        else:
            # 图片文件
            st.info("📷 检测到图片文件")
            st.image(uploaded_file, caption="上传的试卷", use_container_width=True)

            if st.button("🚀 开始分析", type="primary", use_container_width=True):
                with st.status("🔍 AI分析中...", expanded=True) as status:
                    st.write("1. 处理图像...")
                    image = Image.open(io.BytesIO(file_bytes))
                    enhanced = enhance_image_for_ocr(image)

                    st.write("2. AI分析...")
                    try:
                        genai.configure(api_key=GEMINI_API_KEY)
                        model = genai.GenerativeModel(PRIMARY_MODEL)

                        prompt = get_deep_diagnosis_prompt(student_name, exam_type)
                        response = call_ai_with_retry(model, prompt, [enhanced])
                        full_text = response.text

                        st.write("3. 生成雷达图...")

                        if "---JSON_START---" in full_text:
                            parts = full_text.split("---JSON_START---")
                            body = parts[0].strip()
                            json_str = parts[1].strip().replace("```json", "").replace("```", "").strip()
                        else:
                            body = full_text
                            json_str = '{"代数运算": 60, "几何直观": 60, "逻辑推理": 60, "数据分析": 60, "数学建模": 60, "创新意识": 60}'

                        try:
                            radar_data = json.loads(json_str)
                        except:
                            radar_data = {"代数运算": 60, "几何直观": 60, "逻辑推理": 60, "数据分析": 60, "数学建模": 60, "创新意识": 60}

                        st.session_state['report_text'] = body
                        st.session_state['radar_img'] = create_radar_chart_image(radar_data)
                        st.session_state['student_name'] = student_name

                        st.toast("✅ 分析完成！", icon="🎉")
                        status.update(label="✅ 分析完成！", state="complete")
                        st.rerun()

                    except Exception as e:
                        st.error(f"分析失败：{e}")

    # 显示深度报告
    if 'report_text' in st.session_state:
        col1, col2 = st.columns([3, 2])

        with col1:
            st.markdown("<h3 style='color: #ffffff; margin-bottom: 15px;'>深度分析报告</h3>", unsafe_allow_html=True)
            report_html = st.session_state['report_text']

            # 高亮显示分数信息
            report_html = re.sub(r'(\*\*试卷得分.*?\*\*)', r'<span style="color: #FF4B4B;">\1</span>', report_html)
            report_html = re.sub(r'(\*\*正确率.*?\*\*)', r'<span style="color: #FF4B4B;">\1</span>', report_html)

            st.markdown(f"<div class='glass-card'>{report_html}</div>", unsafe_allow_html=True)

        with col2:
            if 'radar_img' in st.session_state:
                st.image(st.session_state['radar_img'], use_container_width=True, caption="能力维度分析")

            docx_file = create_word_docx_simple(
                st.session_state['report_text'],
                st.session_state.get('student_name', '同学'),
                st.session_state.get('radar_img')
            )

            st.download_button(
                "📥 下载完整报告",
                data=docx_file,
                file_name=f"{st.session_state.get('student_name', '同学')}_深度诊断报告.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )

# ==================== 底部信息 ====================
st.markdown("""
<div style='text-align: center; color: #8892b0; font-size: 0.85rem; margin-top: 50px; padding: 20px; border-top: 1px solid #2d3548;'>
    <p>DSE/高考数学诊断工具 v2.1 Pro | 陈老师开发</p>
    <p>AI模型：Gemini 2.5 Pro + GLM-4 Plus 双引擎</p>
</div>
""", unsafe_allow_html=True)
